"""Reusable helpers to build a clean, client-facing Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5F, 0x6B, 0x7A)
ACCENT = RGBColor(0x0B, 0x4F, 0x8A)
RULE = "D0D7DE"


def new_doc():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.9)
        s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    st.paragraph_format.space_after = Pt(7)
    st.paragraph_format.line_spacing = 1.15

    for name, size, color, bold, before, after in [
        ("Title", 26, INK, True, 0, 6),
        ("Heading 1", 17, ACCENT, True, 20, 7),
        ("Heading 2", 13, INK, True, 13, 5),
        ("Heading 3", 11, INK, True, 10, 3),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = bold
        s.font.italic = False
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    return doc


def _shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def _borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), RULE)
        borders.append(e)
    tblPr.append(borders)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _borders(t)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        _shade(c, "F2F5F7")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            c.text = ""
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            _rich(p, str(val))
            for r in p.runs:
                r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def bullets(doc, items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(3)
        _rich(p, it)


def _rich(p, text):
    """Supports **bold** and `code` inline markers."""
    import re
    for tok in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = p.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            p.add_run(tok)


def para(doc, text, italic=False, muted=False):
    p = doc.add_paragraph()
    _rich(p, text)
    for r in p.runs:
        if italic:
            r.italic = True
        if muted:
            r.font.color.rgb = MUTED
    return p


def callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8"); left.set(qn("w:color"), "0B4F8A")
    pbdr.append(left)
    pPr.append(pbdr)
    _rich(p, text)
    for r in p.runs:
        r.font.size = Pt(10)
    return p
