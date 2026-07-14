#!/usr/bin/env python3
"""
Builds the client-facing IAM Engine overview as a Word document.

    pip install python-docx
    python3 docs/client/build_overview_docx.py [output.docx]

Content is deliberately conservative: it describes only what is implemented.
Anything not yet shipped is listed under "Roadmap" and labelled as such.
"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from _docx_style import new_doc, table, bullets, para, callout, MUTED

OUT = sys.argv[1] if len(sys.argv) > 1 else "Coretelligent-IAM-Engine-Overview.docx"
VERSION = "Version 1.0 — 14 July 2026"

doc = new_doc()
H = doc.add_heading

# ─────────────────────────────────────────────────────────── cover
t = doc.add_paragraph("Coretelligent", style="Normal")
t.runs[0].font.size = Pt(12); t.runs[0].bold = True; t.runs[0].font.color.rgb = MUTED
t.paragraph_format.space_after = Pt(2)

ttl = doc.add_paragraph("IAM Engine", style="Title")
ttl.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph("Automated user onboarding and offboarding across your identity and SaaS estate")
sub.runs[0].font.size = Pt(13); sub.runs[0].font.color.rgb = MUTED
sub.paragraph_format.space_after = Pt(14)

v = doc.add_paragraph(VERSION + "  ·  Prepared for client review")
v.runs[0].font.size = Pt(9); v.runs[0].font.color.rgb = MUTED

doc.add_paragraph()
callout(doc, "**In one sentence.** The IAM Engine turns your written onboarding and offboarding runbook into "
             "data, then executes it — every system, every step, in the right order, with an auditable record "
             "of what was done, by whom, and with what result. Credentials are never stored in the platform; "
             "they are held in a vault and released to a single job, at the moment of execution, for the one "
             "system that job is authorized to touch.")

# ─────────────────────────────────────────────────────────── contents
H("What this document covers", 2)
bullets(doc, [
    "**What the platform does** — the problem it solves and the shape of the solution.",
    "**How it works** — how a request becomes a plan, and a plan becomes executed steps.",
    "**The execution paths** — cloud systems, on-premises systems, browser-driven systems, and manual steps.",
    "**The onboarding and offboarding paths** — what actually happens, step by step.",
    "**What you need to configure** — exactly what we need from you, per system.",
    "**Security** — credentials, certificates, agent authentication, access control, approvals, and audit.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 1
H("1. What the platform does", 1)

para(doc, "Every organization has a runbook for bringing a person on and taking a person off. It names the "
          "systems, the groups, the licenses, the mailbox rules, the order things must happen in, and the "
          "things that must never be done without a second pair of eyes. In most organizations that runbook "
          "lives in a document, and it is executed by a person, by hand, under time pressure — usually across "
          "eight to twenty different consoles.")

para(doc, "That is slow, and it is inconsistent, and it is where offboarding misses happen. A forgotten group "
          "membership, a mailbox that was never converted, a VPN account that stays live for six months, an MFA "
          "factor still registered to a phone that walked out of the building.")

para(doc, "The IAM Engine takes that runbook and makes it **executable**. Your runbook becomes structured "
          "configuration — which systems you use, which groups a role gets, what happens on the way out, what "
          "requires approval. A central application reads that configuration, plans the work, and executes each "
          "step against the real system, in dependency order, checking state before it changes it, and recording "
          "everything.")

H("What you get", 2)
bullets(doc, [
    "**Consistency.** The same runbook executes the same way every time, whoever raises the ticket.",
    "**Speed.** A new hire's identity, mailbox, licenses, groups and SaaS access are provisioned in minutes rather than across a day.",
    "**Complete offboarding.** Access is removed everywhere it was granted — including the places a human checklist tends to forget — and the evidence is captured before it is removed.",
    "**An audit trail.** Every action is recorded with the actor, the target, the result, and the evidence. Every step also writes a work note back to the originating ServiceNow ticket.",
    "**Human control where it matters.** Anything destructive is withheld from automation until a named, sufficiently senior person approves it.",
])

H("What it does not do", 2)
para(doc, "The engine is deliberately narrow. It does not replace your identity provider, your MDM, or your "
          "ticketing system — it drives them. It does not hold your secrets; a vault does. It does not make "
          "policy decisions; your runbook does. And it does not take irreversible action on its own.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 2
H("2. How it works", 1)

para(doc, "There are three moving parts.")

table(doc,
      ["Part", "What it is", "Where it lives"],
      [
          ["The application", "The brain. Holds the client configuration, plans each request into steps, queues those "
                              "steps as jobs, brokers credentials, records results, and presents the UI.",
           "Coretelligent-hosted"],
          ["Runners", "The hands. A PowerShell 7 service that picks up a job, executes it against the target system, "
                      "and posts the result back. Same software everywhere; only its placement differs.",
           "Coretelligent cloud, and/or inside your network"],
          ["The vault", "Delinea Secret Server. Holds every credential. The application stores only a reference to a "
                        "secret — never a secret value.",
           "Coretelligent-managed vault"],
      ],
      widths=[1.25, 3.55, 1.7])

H("The lifecycle of a request", 2)

para(doc, "A request moves through six stages. A request is called a **case**; each system a case has to touch is "
          "a **step**.")

table(doc,
      ["Stage", "What happens"],
      [
          ["1. Intake", "A user-management ticket is raised in ServiceNow. The engine reads it — either by polling "
                        "ServiceNow on a schedule, or on demand — and creates a case. Cases are matched to the "
                        "ticket number, so re-reading a ticket never creates a duplicate."],
          ["2. Planning", "The engine loads your profile — the systems you actually use, the groups that go with the "
                          "person's role, what depends on what — and produces an ordered list of steps. Systems you "
                          "do not use are not planned. Systems that only apply to certain roles are only planned when "
                          "that role is matched."],
          ["3. Hold", "Nothing runs on import. A case sits held until an operator releases it, or until its scheduled "
                      "time arrives. An offboard with a termination date is automatically scheduled to release five "
                      "minutes after that instant."],
          ["4. Dispatch", "Each step becomes a job. Jobs are handed out only when their dependencies have succeeded, "
                          "only to a runner that is capable of executing them, and — if the step is destructive — only "
                          "after it has been approved."],
          ["5. Execution", "A runner claims the job, requests the one credential that job needs, connects to the "
                           "target system, checks the current state, makes the change if it is needed, and reads the "
                           "result back to confirm."],
          ["6. Record", "The result is posted back. The engine writes an audit entry, appends to the permanent run "
                        "history, posts a work note to the ServiceNow ticket, and releases whatever steps were "
                        "waiting on that one. When every step is done, the case closes."],
      ],
      widths=[1.15, 5.35])

H("Order is enforced, not assumed", 2)
para(doc, "Steps declare what they depend on, and the engine sorts them. For a client with on-premises Active "
          "Directory, the identity chain is fixed and cannot be misconfigured into a deadlock: "
          "**Active Directory → directory sync → Entra ID → Microsoft 365 → Exchange**. Everything else hangs off "
          "that. Mailbox-dependent steps do not run before the mailbox exists. Licence-dependent steps do not run "
          "before the licence is assigned.")

H("Every step is idempotent", 2)
para(doc, "Every executor checks state before it changes it. If a step is re-run after a partial failure — or run "
          "twice by accident — it converges on the same end state rather than creating a duplicate or a conflict. "
          "This is what makes it safe to retry, and it is the property that lets the engine recover automatically "
          "when a runner dies mid-job.")

H("Dry run", 2)
para(doc, "Any case can be run in **dry-run** mode. Every step connects for real, reads for real, and reports "
          "exactly what it *would* change — and writes nothing. This is the recommended first run for a newly "
          "configured client. Dry-run is enforced at the point a job is handed to a runner, so a case cannot be "
          "half-switched out of it.")

H("Verification, not assumption", 2)
para(doc, "A step reporting success is not the end. Before a case is allowed to complete, the engine re-runs each "
          "step's read-only validator and confirms the intended end state is actually present in the target system. "
          "For onboarding into a hybrid environment it also verifies the source anchor linking the AD account to the "
          "cloud account matches — which is what stops a rehire from silently becoming a second, duplicate identity.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 3
H("3. The execution paths", 1)

para(doc, "Systems differ in where they can be reached from, and in whether they have an API worth using. The "
          "engine has four paths, and picks the right one per step. All four are visible on the case — nothing is "
          "silently skipped.")

H("Path A — Cloud systems, executed centrally", 2)
para(doc, "Microsoft 365, Entra ID, Exchange Online, Google Workspace, Mimecast, Adobe, Zoom, Spanning, KnowBe4, "
          "Egnyte, SentinelOne, Duo, Jira, HubSpot, Salesforce, LogicMonitor, xMatters, Proofpoint and 1Password are "
          "reachable over the public internet. A runner in Coretelligent's environment calls their APIs directly, "
          "authenticating as an application registration or API principal that **you create and control**.")
para(doc, "**Nothing is installed in your network for these systems.** If you are a cloud-only organization, that is "
          "the entire footprint: some application registrations in your tenant, and nothing else.")

H("Path B — On-premises systems, executed by an agent in your network", 2)
para(doc, "Active Directory, directory sync, and hybrid Exchange cannot be reached from outside your network, and "
          "we do not ask you to make them reachable. Instead, a lightweight agent — the same runner software — is "
          "installed on a domain-joined Windows host inside your network.")

callout(doc, "**The agent makes outbound connections only.** It polls the application over HTTPS, asks \"is there "
             "work for me?\", does the work, and posts the result back. The platform never dials into your network. "
             "There is no listening port on the agent, and **no inbound firewall rule is required**. This is the same "
             "pattern as a ServiceNow MID Server.")

para(doc, "What the agent needs:")
bullets(doc, [
    "A **domain-joined Windows host** — a management or jump host is preferred; a domain controller works.",
    "**PowerShell 7** and the **RSAT Active Directory** module. The installer will attempt to add RSAT itself, and will tell you plainly if it cannot.",
    "**Outbound HTTPS** to the application endpoint. Nothing else.",
])
para(doc, "The agent installs as a Windows Scheduled Task running as SYSTEM, restarts on failure, and updates itself "
          "(see below). If you also want your Microsoft 365 work to run from inside your network rather than from "
          "Coretelligent's cloud runner, that is a single configuration flag — all jobs for your organization will "
          "then be claimed by your own agent.")

H("Redundancy across domain controllers", 3)
para(doc, "You can install more than one agent. Each is given a priority. The lower-priority agent stays dormant "
          "while a higher-priority peer is healthy, and takes over automatically when that peer stops reporting in. "
          "Two agents at equal priority share the load instead. No configuration change is needed at failover.")

H("Self-update", 3)
para(doc, "Agents update themselves. The application publishes the current runner build; an agent that is behind "
          "pulls the changed files and restarts itself. An out-of-date agent is not given work — it is brought "
          "current first. Updates can be automatic (the default) or operator-triggered. You do not have to schedule "
          "maintenance windows for the agent, and we do not need remote access to your host to patch it.")

H("Path C — Browser automation, as a last resort", 2)
para(doc, "A small number of systems have no API for the thing that needs doing. For those — and only those — the "
          "runner drives a headless browser through the vendor's own admin console, exactly as a person would.")
bullets(doc, [
    "It runs **inside the runner**, on the same host that already has reach — it is an executor, not a transport.",
    "Passwords are passed to the browser process on standard input only. They are never written to a log, a command line, or a temporary file.",
    "Where the console requires a second factor, the one-time code is minted **by the vault at the moment the prompt appears** — the authenticator seed never leaves the vault. Push, SMS and phone-call factors cannot be automated and the flow stops cleanly rather than guessing.",
    "A browser step that fails is recorded as a **warning with a screenshot**, and does not fail the case.",
])
para(doc, "Today exactly one browser flow ships: forcing a Spanning directory sync, because Spanning's API has no "
          "endpoint for it. We treat browser automation as a liability to be retired, not a strategy — every one is "
          "replaced with an API call the moment the vendor offers one.")

H("Path D — Manual steps, as first-class checklist items", 2)
para(doc, "Some things are not automatable and should not pretend to be: shipping a laptop, the welcome call, "
          "collecting returned equipment, a physical address-book update. These are planned into the case as "
          "**checklist items**. The case will not report itself complete while any of them is outstanding. An "
          "operator ticks them off, and can un-tick them. They are never silently skipped, and they are never quietly "
          "dropped from the plan.")

H("Capability-aware routing", 2)
para(doc, "Each runner reports what it is actually able to do — whether the Active Directory tooling loaded, whether "
          "the browser components are present. The application will not hand an Active Directory job to a runner that "
          "cannot execute one. The job **waits, with a stated reason**, rather than being dispatched to fail. This is "
          "the difference between \"your onboarding is blocked because the agent is missing RSAT\" and a red error at "
          "two in the morning.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 4
H("4. The onboarding path", 1)

para(doc, "A typical onboarding plan for an organization with on-premises Active Directory synchronized to "
          "Microsoft 365. A cloud-only organization simply has no AD or sync links in the chain; the rest is the same.")

table(doc,
      ["#", "Step", "What happens"],
      [
          ["1", "ServiceNow", "The contact and task records are established against the case."],
          ["2", "Active Directory", "The user is created in the correct OU using your username convention. Display "
                                    "name, given name, manager, department, company, office, telephone and proxy "
                                    "addresses are set. Home drive is mapped. Security groups are added — both the "
                                    "baseline set and the conditional ones your rules attach to the role, location or "
                                    "department."],
          ["3", "Directory sync", "A delta sync cycle is triggered and confirmed, so the account appears in the cloud "
                                  "before anything tries to license it."],
          ["4", "Entra ID / Microsoft 365", "The cloud account is confirmed, licensed (directly or via a licensing "
                                            "group), added to cloud groups, given aliases, and — where configured — "
                                            "issued a Temporary Access Pass so the user can register MFA on day one "
                                            "without a shared password."],
          ["5", "Exchange", "Mailbox is enabled and configured. On a hybrid tenant, the mailbox is enabled "
                            "on-premises, because a cloud-side change would simply be overwritten by directory sync."],
          ["6", "Write-back and consistency check", "The assigned email address is written back into Active Directory, "
                                                    "and the source anchor is verified to match the cloud object — so "
                                                    "a rehire re-links to the existing identity rather than creating a "
                                                    "duplicate."],
          ["7", "SaaS estate", "Mimecast, Adobe, Zoom, KnowBe4, Spanning, Egnyte, SharePoint, MDM enrolment, phone "
                               "system, print management, and whatever else your profile lists — each with the groups, "
                               "licenses and product profiles your runbook specifies."],
          ["8", "Manual items", "Workstation build, welcome letter, first-day call — the checklist."],
          ["9", "Case resolution", "Credentials are delivered, MFA registration is confirmed, tasks are closed. This "
                                   "step always runs last."],
      ],
      widths=[0.3, 1.35, 4.85])

H("Password and credential delivery", 2)
para(doc, "The initial password is generated at dispatch. It is shown to an operator **exactly once**, and is wiped "
          "at the moment it is revealed — two people opening the case cannot both see it; the second is told plainly "
          "that it has already been revealed and cannot be recalled. The value is never written to the run log, the "
          "audit record, or the ServiceNow work note. The audit records *that* it was revealed, and by whom — never "
          "*what* it was.")
para(doc, "Where your tenant supports it, we prefer to issue no password at all: a **Temporary Access Pass** lets the "
          "new starter register their own credentials and MFA directly, and nothing reusable ever transits a person.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 5
H("5. The offboarding path", 1)

para(doc, "Offboarding is designed around one principle: **contain first, destroy later, and never destroy without a "
          "human saying so.** Access removal is immediate and reversible. Anything irreversible is gated.")

table(doc,
      ["#", "Step", "What happens"],
      [
          ["1", "Capture evidence", "Before anything is removed, the user's current state is captured and attached to "
                                    "the case — every group membership, every application assignment. If the "
                                    "termination is disputed, or the person is reinstated, the record of what they had "
                                    "is on the case."],
          ["2", "Active Directory", "Password is reset (and captured for the manager, where your runbook says so). All "
                                    "group memberships are removed. The user is hidden from the address book, the "
                                    "manager link is cleared, the account is disabled, and — unless your profile "
                                    "carries the do-not-move guardrail — the object is moved to the Disabled Users OU."],
          ["3", "Entra ID", "The account is confirmed disabled, cloud group memberships and enterprise-application "
                            "assignments are removed, **registered MFA factors are stripped**, and active sessions are "
                            "revoked."],
          ["4", "Exchange", "Mailbox is converted to shared, or forwarded, or given an out-of-office and a delegate — "
                            "whatever your runbook specifies. Delegated access is granted to the named recipient."],
          ["5", "Endpoint", "Where SentinelOne is in scope, the departing user's registered devices are identified and "
                            "disconnected from the network. Isolation is reversible; shutdown is not, and is gated."],
          ["6", "SaaS estate", "Access removed and seats reclaimed across the estate — Mimecast, Adobe, Zoom, Spanning, "
                               "Duo, VPN, Jira, and the rest. Licence downticks happen after the mailbox conversion, "
                               "not before."],
          ["7", "Data custody", "Drive and file ownership transfer, per your runbook."],
          ["8", "Deferred archive", "Where a grace period applies (typically 30–90 days), the archive/delete step is "
                                    "scheduled rather than executed. An immediate-termination flag collapses the grace "
                                    "period to now."],
          ["9", "Equipment return", "Checklist item."],
      ],
      widths=[0.3, 1.2, 5.0])

H("Approval gates", 2)
para(doc, "Each step is classified by intent. **Containment** steps — disable, remove groups, revoke sessions — are "
          "reversible and run on the normal path. **Destructive** steps — deleting a mailbox, hard-deleting an "
          "identity, shutting down a device — are irreversible.")

callout(doc, "A step classified as destructive **always** requires approval and **always** captures evidence, and "
             "that cannot be switched off by configuration. The gate is enforced where the job is handed to a runner, "
             "not in the user interface — an unapproved destructive job is never given to a runner at all, no matter "
             "what any screen does.")

para(doc, "Approval is a separate permission held by senior roles. **An engineer who runs the case cannot approve its "
          "destructive steps** — that requires an operations manager or above. The approver's identity is recorded "
          "against the job automatically; it is not a free-text field.")

H("Guardrails", 2)
para(doc, "Client-specific hazards are encoded as guardrails in your profile rather than lived in someone's memory. "
          "Examples in production today:")
bullets(doc, [
    "**do-not-move-ou** — for tenants where moving the AD object to a Disabled OU takes it out of sync scope and deletes the cloud user. The step disables in place instead.",
    "**do-not-delete** — the identity is disabled and retained, never removed.",
    "**no-device-wipe-without-approval** — endpoint destruction is always gated.",
])

H("Scheduled offboards", 2)
para(doc, "Where the ticket carries a real termination timestamp, the case is automatically scheduled to release "
          "five minutes after it. The engine deliberately **refuses to auto-schedule** — and holds the case for a "
          "human — when the date is ambiguous (a date with no time), already in the past (a backdated ticket must "
          "never fire an unwatched destructive run), or implausibly far out (a mis-keyed year). An offboard whose "
          "target identity could not be resolved with confidence never runs unattended.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 6
H("6. What we need from you", 1)

para(doc, "This section is the practical one. For each system in scope we need a **service principal you create in "
          "your own tenant**, scoped to the minimum permissions the automation actually uses. You retain ownership "
          "of every one of them, and you can revoke any of them at any time, unilaterally, without our involvement.")

H("How credentials are handled — before we list them", 2)
bullets(doc, [
    "You create the credential in **your** system. You hold it.",
    "The credential value is placed in **Delinea Secret Server**. It is never emailed, never pasted into a ticket, and never stored in the application.",
    "The application database holds a **reference** — a vault secret ID — and nothing else. There is no field in the platform that can hold a secret value.",
    "At execution time, the application resolves that one secret and pushes it down to the one runner executing the one job that needs it. See section 7.",
])

H("The setup wizard", 2)
para(doc, "Configuration is not a spreadsheet exchange. Each system has a guided setup page in the application that "
          "walks through the vendor's own console, names the exact screens and the exact permissions, and then "
          "verifies the result in five stages:")

table(doc,
      ["Stage", "What it proves"],
      [
          ["Wired", "A vault reference exists for this system."],
          ["Field check", "The secret actually carries the fields this connector reads — before we try to use it."],
          ["Connection test", "A runner resolved the secret, connected to the live system, and performed one cheap authorized read."],
          ["Rights probe", "Each individual operation the automation will perform is probed and reported — create a user, add to a group, read licences. Where a vendor exposes no way to introspect permissions, we say so rather than guess."],
          ["Dry run", "The real case, executed read-only against the real system, showing exactly what it would change."],
      ],
      widths=[1.15, 5.35])

para(doc, "A system that you handle by hand can be explicitly marked **not needed**. It is then shown as a checklist "
          "item, never as a failure.", muted=True)

# ── 6.1 M365
H("6.1  Microsoft 365 and Entra ID", 2)
para(doc, "This is the one that matters most, and the one with the most nuance.")

H("You create: an Entra application registration", 3)
bullets(doc, [
    "**Single tenant. No redirect URI.** It is an unattended service principal, not a sign-in app.",
    "A **client secret** (24-month expiry). We surface its expiry date in the UI and warn you before it lapses.",
    "The **Directory (tenant) ID**.",
])

callout(doc, "**A Global Administrator user account cannot be used, and never will be able to be.** Entra rejects a "
             "user account in the client-credentials flow — it fails with AADSTS700016 (\"no application with that "
             "app id exists in this tenant\") because the app ID it is being handed is a person, not an application. "
             "This is not a configuration we can work around; it is how the protocol works. The platform actively "
             "detects this and refuses to accept a user account in an application-registration slot.")

H("Microsoft Graph — application permissions, admin-consented", 3)
table(doc,
      ["Permission", "Why", "Required?"],
      [
          ["User.ReadWrite.All", "Create and update users; assign licences.", "Yes"],
          ["Group.ReadWrite.All", "Add and remove group memberships. (GroupMember.ReadWrite.All is sufficient if you prefer it.)", "Yes"],
          ["Organization.Read.All", "Read licence and seat counts, so we can warn before a case fails for want of a seat.", "Yes"],
          ["Domain.Read.All", "Read verified email domains — required for tenants with more than one.", "Yes"],
          ["UserAuthenticationMethod.ReadWrite.All", "Offboarding: strip the leaver's registered MFA factors (phone, Authenticator, FIDO2). Without it, a departed user's registered factors remain, and the engine raises a warning rather than failing.", "Strongly recommended"],
          ["Directory.Read.All", "Resolve managers by name.", "If you set managers"],
          ["Application.Read.All", "Lets the app read its own credential expiry, so we can warn you before your client secret lapses.", "Optional"],
          ["Exchange.ManageAsApp", "Exchange Online administration. Office 365 Exchange Online API, not Graph.", "Only if Exchange is in scope"],
      ],
      widths=[1.85, 3.85, 0.8])

para(doc, "By design, **the application registration is not granted permission to grant itself permissions** — it "
          "holds neither Application.ReadWrite.All nor AppRoleAssignment.ReadWrite.All. Adding a Graph permission is "
          "always a deliberate act by one of your administrators. This is a constraint we impose on ourselves.")

H("Exchange Online — certificate, not secret", 3)
para(doc, "Exchange Online's app-only authentication does not accept a client secret. It requires a **certificate**.")
bullets(doc, [
    "You upload the **public certificate (.cer)** to the same app registration. A self-signed certificate is fine.",
    "The **private key (.pfx)** goes into the vault — base64-encoded, with its password — and never onto disk in the application.",
    "The app registration's service principal must additionally be assigned the **Exchange Administrator** directory role, and it must be **active, not PIM-eligible**. The Exchange.ManageAsApp permission alone is not sufficient — this catches people out.",
])
para(doc, "At execution, the runner writes the private key to a randomly named temporary file only for as long as the "
          "Exchange connection is being established, and deletes it in a guaranteed cleanup block. It is never "
          "committed to a certificate store on a Coretelligent host, and never persisted.")
para(doc, "If you run your own agent, you can instead install the certificate into that host's Windows certificate "
          "store and reference it by thumbprint — the private key then never leaves your network at all.", muted=True)

H("Temporary Access Pass", 3)
para(doc, "If you want new starters to self-register credentials and MFA rather than receive a password, enable "
          "**Temporary Access Pass** in Entra → Protection → Authentication methods, and target it at the relevant "
          "users. This uses the same application registration.")

# ── 6.2 AD
H("6.2  Active Directory (and hybrid Exchange)", 2)
para(doc, "Required only if you have on-premises AD. Executed by an agent in your network — see section 3, Path B.")

H("You provide: a host", 3)
bullets(doc, [
    "A **domain-joined Windows host** with PowerShell 7 and the RSAT Active Directory module. Outbound HTTPS only; no inbound rules.",
])

H("You provide: AD rights (optionally, a service account)", 3)
para(doc, "The agent can run under an account that already holds the necessary rights — in which case **no AD "
          "credential is stored at all**, and it simply uses the ambient domain context. If you prefer a dedicated "
          "service account, it needs:")
bullets(doc, [
    "**Create user objects** delegated on the target OU(s) — this is genuinely probed. We read the OU's security descriptor and tell you, by name, if no access control entry grants the account that right, rather than letting you discover it on the first live onboard.",
    "Modify, disable and move rights on those OUs, for the offboard path.",
])

H("Directory sync", 3)
para(doc, "The sync cmdlets exist only on the Entra Connect server, which is frequently not a domain controller. Name "
          "that host in your profile and the agent will remote into it using the same AD credential; that account must "
          "be permitted to run a sync cycle there.")

H("Hybrid Exchange", 3)
para(doc, "If you run hybrid Exchange, mailbox enablement and conversion **must** happen on-premises — a cloud-side "
          "change is simply overwritten by the next directory sync. We need an AD account with **Exchange Recipient "
          "Management** rights (frequently the same account as above) and your Exchange PowerShell endpoint URI. Note "
          "that this must be the **internal FQDN** the service principal name actually matches, not your public mail "
          "domain — a mismatch here is the single most common hybrid setup failure we see.")

callout(doc, "**A note on OU paths.** \"The server is unwilling to process the request\" on user creation is, in our "
             "experience, almost never a permissions problem — it is a wrong distinguished name. The two causes are a "
             "domain mismatch (the OU path was built from the email domain rather than the AD domain — your AD may be "
             "corp.example.com while your mail is example.com) and OU name spacing. The engine now derives distinguished "
             "names from the domain itself rather than from the email domain, which eliminates the first class entirely.")

# ── 6.3 Google
H("6.3  Google Workspace", 2)
para(doc, "You provide two artifacts:")
bullets(doc, [
    "A **Google Cloud service account with a downloaded JSON key**, in a project with the **Admin SDK API** enabled. No project IAM roles are needed.",
    "A **domain-wide delegation** authorization for that service account's client ID, added in the Admin console under Security → Access and data control → API controls → Manage Domain-Wide Delegation.",
    "A **super-administrator email address** for the service account to impersonate.",
])
para(doc, "The exact OAuth scopes to authorize:")
table(doc,
      ["Scope", "Why"],
      [
          ["https://www.googleapis.com/auth/admin.directory.user", "Create, update and suspend users."],
          ["https://www.googleapis.com/auth/admin.directory.group", "Group membership."],
          ["https://www.googleapis.com/auth/admin.directory.orgunit", "Place users in the correct organizational unit."],
          ["https://www.googleapis.com/auth/admin.directory.user.security",
           "Offboarding: sign the leaver out everywhere. This one matters — suspending an account blocks new sign-ins "
           "but does not invalidate tokens already issued, so without this scope a departing user's phone keeps "
           "syncing mail."],
      ],
      widths=[2.6, 3.9])
para(doc, "Google's delegation is all-or-nothing per token request, which has a pleasant consequence: if the "
          "connection test passes, every requested scope is provably granted. There is no partial state to discover "
          "later. Onboarding will refuse to place a user in the root OU; offboarding suspends and moves to an inactive "
          "OU, and never deletes.", muted=True)

# ── 6.4 the table
H("6.4  The rest of the estate", 2)
para(doc, "Each of these is set up once, and only if you use it. The application's setup guide for each names the "
          "exact console screens.")

table(doc,
      ["System", "Auth method", "What you create and hand over"],
      [
          ["Mimecast", "API 2.0 app, OAuth2 client credentials",
           "An API 2.0 application (Integrations → API and Platform Integrations). Role: Basic Administrator or Help "
           "Desk Administrator, with four products enabled — Account Management, Domain Management, Directory (Sync) "
           "Management, and User & Group Management. Client ID + Client Secret."],
          ["Proofpoint Essentials", "Admin email + password (admin-only API)",
           "A dedicated automation admin login, your pod/region, and your org domain. Note: Proofpoint provisions by "
           "directory sync, not by API — this step verifies the user synced in and retries until it does. It does not "
           "create."],
          ["Spanning Backup", "HTTP Basic — login email + API key",
           "Your Spanning admin login email and a generated API Token, plus your region. If you also want the "
           "force-sync browser flow, the same vault entry additionally needs an M365 admin portal login with a "
           "TOTP/app second factor — push and SMS factors cannot be automated."],
          ["Adobe", "OAuth Server-to-Server (UMAPI v2)",
           "Adobe Developer Console → new project → User Management API → OAuth Server-to-Server. Client ID, Client "
           "Secret, and your Organization ID (…@AdobeOrg). Also the exact product profile names — a typo silently "
           "grants nothing."],
          ["Zoom", "Server-to-Server OAuth app",
           "Marketplace → Develop → Build App → Server-to-Server OAuth. Account ID, Client ID, Client Secret. Six user "
           "scopes; eight more only if you provision Zoom Phone. Turn OFF the \"new experience\" toggle on the build "
           "page — with it on, admin calls fail."],
          ["SentinelOne", "API token (service user)",
           "Settings → Users → Service Users → Create. A service user (not a personal login) with a role that can "
           "disconnect and shut down agents. Its API token, plus your management console URL. Offboard-only."],
          ["Duo", "Admin API, HMAC-signed",
           "Applications → Protect an Application → Admin API, with both \"Grant read resources\" and \"Grant write "
           "resources\". Integration key, secret key, API hostname. Offboard-only."],
          ["KnowBe4", "SCIM 2.0 bearer token",
           "SAML SSO must be configured first — KnowBe4's REST API is read-only and cannot create users. Account "
           "Settings → User Management → SCIM → generate a SCIM bearer token. If you already provision KnowBe4 from "
           "Entra SCIM, you do not need this at all."],
          ["Egnyte", "OAuth2 bearer",
           "An admin service account on your tenant, and your subdomain. (The API key is Coretelligent's and is "
           "reusable.)"],
          ["Jira / Atlassian", "HTTP Basic — admin email + API token",
           "An organization/user-access admin, their API token, and your site URL. Note each product granted consumes "
           "a paid seat."],
          ["Salesforce", "Connected App, OAuth 2.0 JWT bearer",
           "App Manager → New Connected App → Enable OAuth → Use digital signatures, with a certificate you upload. "
           "Scopes: api, and refresh_token/offline_access. Permitted Users = \"Admin approved users are "
           "pre-authorized\". Consumer key, integration user, private key. No password is stored."],
          ["HubSpot", "Private app access token",
           "A Super Admin creates Settings → Integrations → Private Apps, with settings.users.read and "
           "settings.users.write. The access token."],
          ["LogicMonitor", "LMv1 token (HMAC-signed)",
           "A service user with a role that can manage users; Manage → API Tokens → Add. Access ID + Access Key, and "
           "your portal subdomain. Offboard-only."],
          ["xMatters", "HTTP Basic — API key + secret",
           "An account with the REST Web Service User role. Developer → API Keys → Create. Key + secret (the secret is "
           "shown once), plus your company URL."],
          ["Perimeter 81 / Harmony SASE", "API key as bearer",
           "An API key from Settings → API. Access is usually group-driven, so this is typically an on-request step."],
          ["1Password", "SCIM (preferred) or CLI",
           "1Password has no application-only API for user management. Preferred: provision from your IdP via the "
           "SCIM bridge — then nothing is handed over at all, and the engine simply manages the group. The CLI path "
           "requires an owner/admin account exempt from MFA and is not recommended."],
      ],
      widths=[1.05, 1.35, 4.1])

H("Not yet built", 2)
para(doc, "These appear in the catalog but have **no executor today** and are planned as manual checklist steps until "
          "they do: SharePoint, Slack, Teams and Teams Phone, Dropbox, Notion, Printix, Azure Virtual Desktop, MDM "
          "(Intune/Jamf/Addigy), bulk data transfer, and mailbox archive. We would rather tell you this than have you "
          "discover it.")

H("Credential rotation", 2)
para(doc, "Your Entra client secret and your Exchange certificate both expire. The engine reads their expiry from the "
          "vault and from the tenant itself, surfaces it on the client's health view, and raises a notification before "
          "it lapses — so rotation is a scheduled task rather than an outage.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 7
H("7. Security", 1)

para(doc, "The engine holds the keys to your identity estate. The design starts from that premise, and the controls "
          "below are the ones that are actually implemented and in force — not aspirations. Where something is "
          "planned rather than shipped, it is listed as such in section 7.8.")

H("7.1  The central principle: the platform stores no secrets", 2)

para(doc, "Every credential lives in **Delinea Secret Server**. The application's own database contains a vault "
          "**reference** — an ID — and a label. It does not contain, and has no schema field capable of containing, a "
          "credential value. The client profiles are the same: they carry references, never values. This is enforced "
          "in the data model, not by convention.")

H("How a credential reaches a runner", 3)
para(doc, "At the moment a runner needs to act, it asks the application for the one credential that job requires. "
          "Before releasing anything, the application checks, in order:")
bullets(doc, [
    "**The job exists, and this runner owns it.** A runner cannot request a credential for a job assigned to a different runner.",
    "**The runner is enabled.**",
    "**The job is actually in progress.** Credentials are not brokered for pending, completed or failed jobs.",
    "**The requested secret is on that job's allowlist.** Each job carries the specific secret names it is permitted to request. A runner cannot ask for an arbitrary secret — it can only ask for the ones its current job legitimately needs.",
])
para(doc, "Only then is the value resolved from the vault and returned — with cache headers that forbid any "
          "intermediary from retaining it. The request is written to the audit log; the value is not.")

callout(doc, "**The runner has no vault credentials of its own.** It cannot talk to Delinea, and it does not know how "
             "to. It receives only what the application pushes down for the specific job in hand. This means an agent "
             "installed on your domain controller — the machine most exposed to your internal network — holds **no "
             "standing access to any credential whatsoever**. Compromising it does not yield the vault. It yields, at "
             "most, the credentials for jobs currently in flight on that host.")

H("In memory only", 3)
para(doc, "The runner holds the credential in process memory for the lifetime of the job and no longer. It is never "
          "written to disk, never written into a profile, never cached, and never re-used for a subsequent job.")

H("7.2  Secrets never reach a log, a ticket, or an error message", 2)

para(doc, "Error text is the classic leak path — a stack trace containing a connection string, pasted into a ticket. "
          "Before any failure text leaves a runner, it passes through a scrubber that removes:")
bullets(doc, [
    "The values of any field whose name suggests a secret (password, secret, key, token, credential, certificate, private, and so on).",
    "**Any value carrying the structural characters of an encoded blob** — slashes, plus signs, equals signs, braces, quotes, whitespace — regardless of what the field is called. A base64 string or a PEM block is never a hostname, so it is scrubbed on shape alone. This is what catches a secret that arrives in an unexpected field.",
    "Generated passwords injected by the application.",
    "**The runner's own API token.**",
])
para(doc, "Usernames and server names are deliberately left visible, because a redacted error is useless for "
          "diagnosis. The scrubbed text is what is persisted to the run record, shown in the UI, and posted to the "
          "ServiceNow work note — all three read from the same scrubbed source.")
para(doc, "There is a second, independent redaction boundary in front of the AI features used for runbook parsing: "
          "vault URLs, passwords, national identifiers, phone numbers and email local parts are stripped before any "
          "text is sent to a language model. Secrets do not cross that boundary.")

H("7.3  Certificates", 2)
table(doc,
      ["Where", "How"],
      [
          ["Exchange Online", "App-only authentication is certificate-based by Microsoft's design — a client secret "
                              "will not work. You hold the certificate; the public half goes on your app "
                              "registration, the private half goes in the vault. At execution the private key is "
                              "materialized only for the duration of the connection handshake and deleted in a "
                              "guaranteed cleanup path. It is never left on disk, and never installed into a "
                              "certificate store on a Coretelligent host."],
          ["Salesforce", "The Connected App uses the JWT bearer flow — the runner signs an assertion with a private "
                         "key held in the vault. No Salesforce password is ever stored."],
          ["Google Workspace", "The service-account key signs a short-lived assertion. No password exists to steal."],
          ["Your own agent", "If you run your own agent, the Exchange certificate can instead be installed into that "
                             "host's Windows certificate store and referenced by thumbprint — the private key then "
                             "never leaves your network at all."],
      ],
      widths=[1.3, 5.2])

H("7.4  Operator access to the application", 2)

para(doc, "**Sign-in.** Coretelligent staff sign in with **Microsoft Entra SSO** (OpenID Connect, authorization code "
          "flow with PKCE), against Coretelligent's own tenant. A local break-glass account exists for the case where "
          "SSO itself is the outage; its password is stored as a salted scrypt hash and verified in constant time, "
          "and its use is audited under its own distinct event.")

para(doc, "**Sessions are not bearer tokens and are not JWTs.** A session is an opaque, high-entropy random value in "
          "an HTTP-only, same-site, secure cookie. The server stores **only its SHA-256 hash** — so a database "
          "disclosure yields nothing replayable. Sessions expire after 12 hours and can be revoked centrally and "
          "immediately.")

para(doc, "**Authorization is permission-based, across six roles.** Permissions — not role names — are checked at "
          "every server-side entry point. The separations that matter:")
bullets(doc, [
    "An **engineer** can plan and run cases, but **cannot approve destructive steps**.",
    "An **auditor** is strictly read-only.",
    "An **importer** can bring cases in but not execute them.",
    "Granting or removing the highest role is restricted to that role — so an administrator cannot promote themselves out of a control.",
])

para(doc, "**Per-client scoping is a server-side boundary, not a UI filter.** Each operator's access to clients is "
          "either all, an explicit allowlist, or all-except-a-denylist; individual clients can additionally be marked "
          "**restricted**, requiring an explicit grant on top. This is enforced in the data queries themselves — and "
          "requests for a client outside an operator's scope return **not found**, not **forbidden**, so the existence "
          "of a client is not disclosed by probing. Hiding a client from a list is not a boundary if the API still "
          "answers a guessed URL; we treat it accordingly.")

para(doc, "**\"View as\" is read-only.** A super-administrator can view the application as another user to reproduce "
          "an access problem. While doing so, **every mutation is refused** — it is a lens, not a login. Impersonation "
          "cannot be nested.")

H("7.5  How runners authenticate", 2)

para(doc, "Two credentials are involved, and they do different jobs.")

table(doc,
      ["Credential", "What it is", "Lifetime"],
      [
          ["Enrollment token", "A short-lived, HMAC-signed token minted in the UI and bound to a specific scope and a "
                               "specific client. It is **self-describing**: a token minted for your organization can "
                               "only ever register an agent for your organization — it cannot be replayed to enroll "
                               "an agent somewhere else.", "One hour, by default"],
          ["Runner API token", "A bearer token the agent presents on every subsequent call. It is baked into the "
                               "installer at generation time and written to the machine environment so the SYSTEM "
                               "service can read it.", "Until rotated"],
      ],
      widths=[1.2, 4.35, 0.95])

para(doc, "**The machine API is fail-closed.** A request to any runner endpoint without a valid token is rejected. If "
          "the token is not configured at all, the endpoints return a service error rather than opening — a missing "
          "configuration cannot become an open door. The endpoints that carry credentials fail closed in **every** "
          "environment including development, because \"no token configured\" must never resolve to \"serve "
          "tenant-administrator credentials to an unauthenticated caller\".")

para(doc, "The handful of endpoints that must be reachable without a token — the installer bootstrap, which a host "
          "with no token yet has to fetch — are an **explicit allowlist, not a path prefix**. That is a deliberate "
          "choice: it means a new credential-carrying endpoint added tomorrow is protected by default, rather than "
          "accidentally exposed because it happened to share a URL prefix.")

para(doc, "Retrieving the runner token in the application is itself a privileged, audited action available only to "
          "senior administrators — and the audit records the retrieval, never the value.")

H("7.6  Network posture", 2)

callout(doc, "**The guarantee that matters: the agent inside your network makes outbound connections only.** It polls "
             "the application, does work, and reports back. It exposes no listening port. The platform holds no "
             "route, no credential, and no mechanism to initiate a connection into your network. **No inbound "
             "firewall rule is required, and none should be created.** If you shut the agent down, we lose the "
             "ability to act — we do not retain a back door.")

para(doc, "All communication with your SaaS systems, with the vault, and with Entra is over HTTPS. The runner "
          "contains no TLS-validation bypass — certificate validation is never disabled, anywhere in the codebase.")

para(doc, "**Current deployment state, stated plainly:** the platform is running as a pilot from a Coretelligent-hosted "
          "endpoint. Agents reaching it across the public internet do so over HTTPS. The production move to Azure "
          "Container Apps — a single stable HTTPS endpoint with a managed certificate, and platform secrets held in "
          "Azure Key Vault via managed identity — is the next scheduled piece of work, and is what every agent will "
          "poll once complete. We would rather set out where we are than imply a posture we have not yet finished "
          "building.")

H("7.7  Approvals, evidence, and audit", 2)

para(doc, "**Approval is enforced at the dispatch gate.** An unapproved destructive job is never handed to a runner. "
          "This is checked where the work is claimed, not in the interface — so the control cannot be bypassed by a "
          "UI defect, a direct API call, or a misconfigured screen. Approval requires a permission that the person "
          "running the case does not, by default, hold.")

para(doc, "**Evidence precedes removal.** On the offboard path, group memberships and application assignments are "
          "captured and attached to the case before they are removed. The record of what someone had survives the "
          "removal of what they had.")

para(doc, "**The audit log is append-only, and it is comprehensive.** Roughly 112 distinct action types are recorded, "
          "each with an actor, a timestamp, the client, the case, the job, and a structured detail payload. Every "
          "sign-in, every failed sign-in, every credential brokerage, every approval, every agent enrollment, every "
          "secret re-wiring, every password reveal. Auditing is designed so that it cannot break the action it "
          "records — but the actions it records cannot escape it.")

para(doc, "Alongside it, a permanent **run history** keeps one immutable row per job result. A job re-run overwrites "
          "the job's current state, but never the history — the record of the first attempt, and what it did, remains.")

para(doc, "**Write-back to ServiceNow.** Every step posts a work note to the originating ticket, so the record of what "
          "was done lives where your service desk already looks. Work-note text passes through the same secret "
          "scrubber. The ticket number is pattern-validated before it is used in a query, which closes the injection "
          "path that a naively constructed ServiceNow query would otherwise open.")

H("7.8  What is planned, and not yet shipped", 2)

para(doc, "We would rather you read this from us than find it later.")

table(doc,
      ["Item", "Status"],
      [
          ["**Mutual TLS between agents and the application.** Each agent would present a client certificate issued at "
           "enrollment, replacing the shared bearer token.",
           "Planned. Today, agents authenticate with an enrollment token bound to your organization, plus a bearer "
           "token on the machine API."],
          ["**Per-agent credentials.** Today the bearer token is shared across the fleet rather than unique per agent.",
           "Planned, alongside mutual TLS. The enrollment half — a token that can only ever register the agent it was "
           "minted for — is already in place."],
          ["**Automated token rotation.** Rotation today is a configuration change plus a re-run of the installer.",
           "Planned."],
          ["**Azure hosting with a managed TLS certificate and Key Vault–held platform secrets.**",
           "Planned; the next scheduled work item."],
      ],
      widths=[3.6, 2.9])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════ 8
H("8. Summary", 1)

para(doc, "The IAM Engine executes your onboarding and offboarding runbook, across your whole estate, the same way "
          "every time. Cloud systems are driven through APIs from Coretelligent's environment, using service "
          "principals **you create and can revoke**. On-premises systems are driven by a lightweight agent inside "
          "your network that makes outbound connections only and requires no inbound firewall change. Every step "
          "checks state before it changes it, verifies the result afterwards, and records what it did.")

para(doc, "No credential is stored in the platform. The vault holds them; the application holds references; a runner "
          "receives exactly one credential, for exactly one job, at the moment of execution, and holds it only in "
          "memory. Anything irreversible is withheld from automation until a named, senior human approves it — and "
          "the evidence of what someone had is captured before it is taken away.")

para(doc, "Questions, and requests for the detailed setup guide for any individual system, to your Coretelligent "
          "engagement contact.", muted=True)

doc.save(OUT)
print(f"wrote {OUT}")
